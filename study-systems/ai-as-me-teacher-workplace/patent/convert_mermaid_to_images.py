#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Mermaid图表转图片并插入Word文档
使用在线API或本地工具转换Mermaid图表
"""

import os
import sys
import re
import subprocess
import base64
import json
from pathlib import Path
import tempfile
import urllib.request
import urllib.parse

# 设置输出编码
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

def check_mermaid_cli():
    """检查mermaid-cli是否安装"""
    try:
        result = subprocess.run(['mmdc', '--version'], capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False

def convert_mermaid_online(mermaid_code, output_file):
    """使用在线API转换Mermaid代码为PNG"""
    try:
        # 使用mermaid.ink API
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.ink/img/{encoded}"
        
        # 下载图片
        urllib.request.urlretrieve(url, output_file)
        return True
    except Exception as e:
        print(f"在线转换失败: {e}")
        # 尝试使用mermaid.live API
        try:
            # 使用mermaid.live的API
            api_url = "https://mermaid.live/api/v1/svg"
            data = json.dumps({"code": mermaid_code}).encode('utf-8')
            req = urllib.request.Request(api_url, data=data, headers={'Content-Type': 'application/json'})
            
            with urllib.request.urlopen(req) as response:
                svg_content = response.read().decode('utf-8')
                
            # 将SVG转换为PNG（需要cairosvg）
            try:
                import cairosvg
                cairosvg.svg2png(bytestring=svg_content.encode('utf-8'), write_to=str(output_file))
                return True
            except ImportError:
                # 如果没有cairosvg，保存为SVG
                svg_file = output_file.with_suffix('.svg')
                svg_file.write_text(svg_content, encoding='utf-8')
                print(f"  已保存为SVG: {svg_file}")
                return False
        except Exception as e2:
            print(f"备用API也失败: {e2}")
            return False

def convert_mermaid_to_png(mermaid_code, output_file, title=""):
    """转换Mermaid代码为PNG"""
    # 优先使用mermaid-cli
    if check_mermaid_cli():
        return convert_mermaid_cli(mermaid_code, output_file)
    else:
        # 使用在线API
        return convert_mermaid_online(mermaid_code, output_file)

def convert_mermaid_cli(mermaid_code, output_file):
    """使用mermaid-cli转换"""
    temp_file = Path(tempfile.gettempdir()) / f"mermaid_{os.getpid()}.mmd"
    
    try:
        temp_file.write_text(mermaid_code, encoding='utf-8')
        
        cmd = [
            'mmdc',
            '-i', str(temp_file),
            '-o', str(output_file),
            '-w', '1920',
            '-H', '1080',
            '-b', 'white',
            '--scale', '2'
        ]
        
        subprocess.run(cmd, capture_output=True, check=True)
        return True
    except Exception as e:
        return False
    finally:
        if temp_file.exists():
            temp_file.unlink()

def extract_mermaid_blocks(md_file):
    """从Markdown文件中提取Mermaid代码块"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 匹配mermaid代码块，包括标题
    pattern = r'##\s*(图\d+[：:].*?)\n\n```mermaid\n(.*?)```'
    matches = re.findall(pattern, content, re.DOTALL)
    
    result = []
    for title, code in matches:
        result.append({
            'title': title.strip(),
            'code': code.strip()
        })
    
    return result

def update_word_with_images(word_file, images_dir, mermaid_blocks):
    """更新Word文档，插入图片"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document(word_file)
        
        # 清除原有内容（保留第一个标题）
        doc.paragraphs.clear()
        doc.add_heading('附图（Mermaid图表）', 0)
        
        # 插入每个图片
        for i, block in enumerate(mermaid_blocks, 1):
            title = block['title']
            figure_match = re.search(r'图(\d+)', title)
            figure_num = figure_match.group(1) if figure_match else str(i)
            
            # 清理标题
            clean_title = title.replace(f'图{figure_num}：', '').replace(f'图{figure_num}:', '').strip()
            
            # 查找图片文件
            image_file = images_dir / f"图{figure_num}.png"
            if not image_file.exists():
                image_file = images_dir / f"图{figure_num}.svg"
            
            if image_file.exists():
                # 添加标题
                doc.add_heading(f"图{figure_num}：{clean_title}", 2)
                
                # 插入图片
                if image_file.suffix == '.png':
                    doc.add_picture(str(image_file), width=Inches(6.5))
                else:
                    # SVG需要转换为PNG
                    print(f"  注意: {image_file.name} 是SVG格式，建议转换为PNG")
                    doc.add_paragraph(f"[图片: {image_file.name}]")
                
                # 添加说明
                para = doc.add_paragraph()
                para.add_run(f"图{figure_num}：{clean_title}").bold = True
                para.alignment = 1  # 居中
                
                doc.add_paragraph()  # 空行
        
        # 保存
        doc.save(word_file)
        return True
    except ImportError:
        print("  [WARN] python-docx未安装，无法自动更新Word文档")
        print("  请手动将图片插入Word文档")
        return False
    except Exception as e:
        print(f"  [FAIL] 更新Word文档失败: {e}")
        return False

def process_patent(patent_dir, patent_name):
    """处理一个专利的Mermaid图表"""
    print(f"\n处理: {patent_name}")
    print("-" * 60)
    
    # 查找Mermaid图表文件
    mermaid_file = None
    for filename in ["06-附图（Mermaid图表）.md", "07-附图（Mermaid图表）.md"]:
        test_file = patent_dir / filename
        if test_file.exists():
            mermaid_file = test_file
            break
    
    if not mermaid_file:
        print(f"未找到Mermaid图表文件")
        return False
    
    # 提取Mermaid代码块
    mermaid_blocks = extract_mermaid_blocks(mermaid_file)
    print(f"找到 {len(mermaid_blocks)} 个图表")
    
    if not mermaid_blocks:
        print("未找到Mermaid代码块")
        return False
    
    # 创建图片输出目录
    images_dir = patent_dir / "word-output" / "images"
    images_dir.mkdir(exist_ok=True)
    
    # 转换每个图表
    success_count = 0
    for i, block in enumerate(mermaid_blocks, 1):
        title = block['title']
        code = block['code']
        
        # 提取图号
        figure_match = re.search(r'图(\d+)', title)
        figure_num = figure_match.group(1) if figure_match else str(i)
        
        # 生成输出文件名
        image_file = images_dir / f"图{figure_num}.png"
        
        print(f"转换图{figure_num}: {title}")
        
        # 转换
        success = convert_mermaid_to_png(code, image_file, title)
        
        if success:
            print(f"  [OK] 图片已生成: {image_file.name}")
            success_count += 1
        else:
            print(f"  [FAIL] 转换失败，尝试其他方法...")
            # 尝试使用在线API
            if convert_mermaid_online(code, image_file):
                print(f"  [OK] 使用在线API转换成功")
                success_count += 1
            else:
                print(f"  [FAIL] 所有方法都失败")
    
    # 更新Word文档
    word_file = None
    for filename in ["06-附图（Mermaid图表）.docx", "07-附图（Mermaid图表）.docx"]:
        test_file = patent_dir / "word-output" / filename
        if test_file.exists():
            word_file = test_file
            break
    
    if word_file:
        print(f"\n更新Word文档: {word_file.name}")
        if update_word_with_images(word_file, images_dir, mermaid_blocks):
            print(f"  [OK] Word文档已更新")
        else:
            print(f"  [WARN] 请手动更新Word文档")
    
    print(f"\n完成: {success_count}/{len(mermaid_blocks)} 个图表转换成功")
    return success_count == len(mermaid_blocks)

def main():
    """主函数"""
    print("=" * 60)
    print("Mermaid图表转图片工具")
    print("=" * 60)
    
    # 检查工具
    has_cli = check_mermaid_cli()
    if has_cli:
        print("使用mermaid-cli进行转换")
    else:
        print("mermaid-cli未安装，使用在线API转换")
        print("(建议安装: npm install -g @mermaid-js/mermaid-cli)")
    
    patent_dir = Path(__file__).parent
    
    # 定义三个专利
    patents = [
        {
            "name": "专利1-多层次记忆架构的AI赋能工作空间系统",
            "dir": patent_dir,
        },
        {
            "name": "专利2-物流数字孪生教学平台",
            "dir": patent_dir / "专利2-物流数字孪生教学平台",
        },
        {
            "name": "专利3-物流实践课程智能设计系统",
            "dir": patent_dir / "专利3-物流实践课程智能设计系统",
        }
    ]
    
    total_success = 0
    total_patents = 0
    
    for patent in patents:
        if patent["dir"].exists():
            total_patents += 1
            if process_patent(patent["dir"], patent["name"]):
                total_success += 1
    
    print()
    print("=" * 60)
    print(f"处理完成！成功: {total_success}/{total_patents} 个专利")
    print("=" * 60)
    
    if total_success < total_patents:
        print("\n提示:")
        print("1. 如果转换失败，可以手动使用 https://mermaid.live/ 转换")
        print("2. 图片已保存在各自的 word-output/images/ 目录")
        print("3. 可以手动将图片插入Word文档")
    
    return 0 if total_success == total_patents else 1

if __name__ == "__main__":
    sys.exit(main())
