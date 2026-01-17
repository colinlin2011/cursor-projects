#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
更新Word文档，插入Mermaid图表图片
"""

import sys
import re
from pathlib import Path

# 设置输出编码
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("错误: python-docx未安装")
    print("请运行: pip install python-docx")
    sys.exit(1)

def extract_figure_info(md_file):
    """从Markdown文件中提取图表信息"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'##\s*(图\d+[：:].*?)\n\n```mermaid\n(.*?)```'
    matches = re.findall(pattern, content, re.DOTALL)
    
    result = []
    for title, code in matches:
        figure_match = re.search(r'图(\d+)', title)
        figure_num = figure_match.group(1) if figure_match else None
        clean_title = title.replace(f'图{figure_num}：', '').replace(f'图{figure_num}:', '').strip() if figure_num else title.strip()
        
        result.append({
            'figure_num': figure_num,
            'title': clean_title,
            'full_title': title.strip()
        })
    
    return result

def update_word_document(word_file, images_dir, figure_info):
    """更新Word文档，插入图片"""
    doc = Document(word_file)
    
    # 清除原有内容
    doc.paragraphs.clear()
    
    # 添加标题
    doc.add_heading('附图（Mermaid图表）', 0)
    doc.add_paragraph()  # 空行
    
    # 插入每个图片
    for info in figure_info:
        figure_num = info['figure_num']
        title = info['title']
        full_title = info['full_title']
        
        # 查找图片文件
        image_file = images_dir / f"图{figure_num}.png"
        if not image_file.exists():
            print(f"  警告: 图片不存在 {image_file.name}")
            continue
        
        # 添加图表标题
        doc.add_heading(f"图{figure_num}：{title}", 2)
        
        # 插入图片
        try:
            doc.add_picture(str(image_file), width=Inches(6.5))
        except Exception as e:
            print(f"  错误: 无法插入图片 {image_file.name}: {e}")
            doc.add_paragraph(f"[图片: {image_file.name}]")
        
        # 添加图片说明
        para = doc.add_paragraph()
        para.add_run(f"图{figure_num}：{title}").bold = True
        para.alignment = 1  # 居中
        
        doc.add_paragraph()  # 空行
    
    # 保存
    doc.save(word_file)
    print(f"  [OK] Word文档已更新: {word_file.name}")

def main():
    patent_dir = Path(__file__).parent
    
    patents = [
        {
            "name": "专利1",
            "dir": patent_dir,
            "mermaid_file": "07-附图（Mermaid图表）.md",
            "word_file": "word-output/附图说明.docx"  # 专利1的Word文件可能在这个位置
        },
        {
            "name": "专利2",
            "dir": patent_dir / "专利2-物流数字孪生教学平台",
            "mermaid_file": "06-附图（Mermaid图表）.md",
            "word_file": "word-output/06-附图（Mermaid图表）.docx"
        },
        {
            "name": "专利3",
            "dir": patent_dir / "专利3-物流实践课程智能设计系统",
            "mermaid_file": "06-附图（Mermaid图表）.md",
            "word_file": "word-output/06-附图（Mermaid图表）.docx"
        }
    ]
    
    print("=" * 60)
    print("更新Word文档，插入Mermaid图表图片")
    print("=" * 60)
    
    for patent in patents:
        if not patent["dir"].exists():
            continue
        
        mermaid_file = patent["dir"] / patent["mermaid_file"]
        word_file = patent["dir"] / patent["word_file"]
        images_dir = patent["dir"] / "word-output" / "images"
        
        if not mermaid_file.exists():
            print(f"\n跳过 {patent['name']}: 未找到Mermaid文件")
            continue
        
        if not word_file.exists():
            print(f"\n跳过 {patent['name']}: 未找到Word文件")
            continue
        
        if not images_dir.exists():
            print(f"\n跳过 {patent['name']}: 未找到图片目录")
            continue
        
        print(f"\n处理: {patent['name']}")
        print("-" * 60)
        
        # 提取图表信息
        figure_info = extract_figure_info(mermaid_file)
        print(f"找到 {len(figure_info)} 个图表")
        
        # 更新Word文档
        update_word_document(word_file, images_dir, figure_info)
    
    print("\n" + "=" * 60)
    print("完成！")
    print("=" * 60)

if __name__ == "__main__":
    main()
