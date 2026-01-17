#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
为专利1创建专门的附图Word文档
"""

import sys
import re
from pathlib import Path

# 设置输出编码
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("错误: python-docx未安装")
    sys.exit(1)

def extract_figure_info(md_file):
    """从Markdown文件中提取图表信息"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'##\s*(图\d+[：:].*?)\n\n```mermaid\n(.*?)```'
    matches = re.findall(pattern, content, re.DOTALL)
    
    result = []
    for title, code in matches:
        figure_match = re.search(r'图(\d+)', title)
        figure_num = figure_match.group(1) if figure_match else None
        clean_title = title.replace(f'图{figure_num}：', '').replace(f'图{figure_num}:', '').strip() if figure_num else title.strip()
        
        result.append({
            'figure_num': figure_num,
            'title': clean_title,
            'full_title': title.strip()
        })
    
    return result

def main():
    patent_dir = Path(__file__).parent
    mermaid_file = patent_dir / "07-附图（Mermaid图表）.md"
    images_dir = patent_dir / "word-output" / "images"
    word_file = patent_dir / "word-output" / "07-附图（Mermaid图表）.docx"
    
    if not mermaid_file.exists():
        print("未找到Mermaid文件")
        return
    
    if not images_dir.exists():
        print("未找到图片目录")
        return
    
    print("创建专利1的附图Word文档...")
    
    # 提取图表信息
    figure_info = extract_figure_info(mermaid_file)
    print(f"找到 {len(figure_info)} 个图表")
    
    # 创建Word文档
    doc = Document()
    doc.add_heading('附图（Mermaid图表）', 0)
    doc.add_paragraph()  # 空行
    
    # 插入每个图片
    for info in figure_info:
        figure_num = info['figure_num']
        title = info['title']
        
        # 查找图片文件
        image_file = images_dir / f"图{figure_num}.png"
        if not image_file.exists():
            print(f"  警告: 图片不存在 {image_file.name}")
            continue
        
        # 添加图表标题
        doc.add_heading(f"图{figure_num}：{title}", 2)
        
        # 插入图片
        try:
            doc.add_picture(str(image_file), width=Inches(6.5))
        except Exception as e:
            print(f"  错误: 无法插入图片 {image_file.name}: {e}")
            doc.add_paragraph(f"[图片: {image_file.name}]")
        
        # 添加图片说明
        para = doc.add_paragraph()
        para.add_run(f"图{figure_num}：{title}").bold = True
        para.alignment = 1  # 居中
        
        doc.add_paragraph()  # 空行
    
    # 保存
    doc.save(word_file)
    print(f"[OK] Word文档已创建: {word_file.name}")

if __name__ == "__main__":
    main()
