"""
Markdown转Word工具
将Markdown文档转换为Word格式，保留表格和格式
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from typing import List, Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MarkdownToWord:
    """Markdown转Word转换器"""
    
    def __init__(self):
        pass
    
    def parse_markdown_table(self, markdown_text: str) -> List[List[str]]:
        """解析Markdown表格"""
        lines = markdown_text.strip().split('\n')
        table_data = []
        
        for line in lines:
            line = line.strip()
            if not line or not line.startswith('|'):
                continue
            
            # 跳过分隔行
            if re.match(r'^\|[\s\-\|:]+\|$', line):
                continue
            
            # 解析表格行
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                table_data.append(cells)
        
        return table_data
    
    def add_table_to_doc(self, doc: Document, table_data: List[List[str]]):
        """将表格添加到Word文档"""
        if not table_data or len(table_data) < 2:
            return
        
        # 创建表格
        num_rows = len(table_data)
        num_cols = len(table_data[0])
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # 填充数据
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                
                # 设置表头格式
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    def convert_to_word(self, markdown_file: str, output_file: Optional[str] = None) -> str:
        """
        将Markdown文件转换为Word
        
        Args:
            markdown_file: Markdown文件路径
            output_file: 输出Word文件路径（可选）
        
        Returns:
            输出文件路径
        """
        markdown_path = Path(markdown_file)
        
        if output_file is None:
            output_file = markdown_path.parent / f"{markdown_path.stem}.docx"
        else:
            output_file = Path(output_file)
        
        # 读取Markdown文件
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        
        # 创建Word文档
        doc = Document()
        
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                
                if text:
                    heading = doc.add_heading(text, level=min(level, 9))
                    i += 1
                    continue
            
            # 处理表格
            if '|' in line and not line.strip().startswith('#'):
                # 收集表格行
                table_lines = [line]
                i += 1
                
                # 跳过分隔行
                if i < len(lines) and re.match(r'^\|[\s\-\|:]+\|$', lines[i].strip()):
                    i += 1
                
                # 收集表格数据行
                while i < len(lines):
                    next_line = lines[i].rstrip()
                    if '|' in next_line and not next_line.strip().startswith('#'):
                        table_lines.append(next_line)
                        i += 1
                    else:
                        break
                
                # 解析并添加表格
                table_text = '\n'.join(table_lines)
                table_data = self.parse_markdown_table(table_text)
                if table_data:
                    self.add_table_to_doc(doc, table_data)
                continue
            
            # 处理普通段落
            if line.strip() and not line.startswith('---'):
                # 跳过空行和分隔线
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                i += 1
            else:
                i += 1
        
        # 保存文档
        doc.save(str(output_file))
        logger.info(f"已创建Word文件: {output_file}")
        return str(output_file)


def convert_clause_matrix_to_word():
    """转换条款对应矩阵为Word"""
    matrix_file = r"p:\Cursor Project\study-systems\autonomous-safety-study\materials\04_工作文档\条款对应矩阵.md"
    output_file = r"p:\Cursor Project\study-systems\autonomous-safety-study\materials\04_工作文档\条款对应矩阵.docx"
    
    converter = MarkdownToWord()
    
    try:
        result = converter.convert_to_word(matrix_file, output_file)
        print(f"✅ 转换成功！")
        print(f"Word文件已保存到: {result}")
        return True
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    convert_clause_matrix_to_word()
