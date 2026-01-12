"""
文档转换工具 - 核心转换模块
支持将PDF、Word、Excel转换为Markdown格式
"""

import os
from pathlib import Path
from typing import Optional, Tuple
import logging

# PDF处理
try:
    import pymupdf  # PyMuPDF (fitz)
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word处理
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Excel处理
try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentConverter:
    """文档转换器"""
    
    def __init__(self):
        self.check_dependencies()
    
    def check_dependencies(self):
        """检查依赖库是否安装"""
        missing = []
        if not PDF_AVAILABLE:
            missing.append("PyMuPDF (pip install pymupdf)")
        if not WORD_AVAILABLE:
            missing.append("python-docx (pip install python-docx)")
        if not EXCEL_AVAILABLE:
            missing.append("pandas (pip install pandas openpyxl)")
        
        if missing:
            logger.warning(f"以下依赖库未安装: {', '.join(missing)}")
    
    def convert_pdf_to_markdown(self, pdf_path: str, output_path: Optional[str] = None) -> Tuple[str, bool]:
        """
        将PDF转换为Markdown
        
        Args:
            pdf_path: PDF文件路径
            output_path: 输出文件路径（可选）
        
        Returns:
            (markdown_content, success)
        """
        if not PDF_AVAILABLE:
            return "错误：PyMuPDF库未安装。请运行: pip install pymupdf", False
        
        try:
            import pymupdf as fitz
            
            doc = fitz.open(pdf_path)
            markdown_content = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                # 添加页面分隔符
                if page_num > 0:
                    markdown_content.append("\n---\n")
                
                markdown_content.append(f"## 第 {page_num + 1} 页\n\n")
                markdown_content.append(text)
                markdown_content.append("\n")
            
            doc.close()
            
            result = "".join(markdown_content)
            
            # 保存到文件
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(result)
                logger.info(f"PDF已转换为Markdown: {output_path}")
            
            return result, True
            
        except Exception as e:
            error_msg = f"PDF转换失败: {str(e)}"
            logger.error(error_msg)
            return error_msg, False
    
    def convert_word_to_markdown(self, word_path: str, output_path: Optional[str] = None) -> Tuple[str, bool]:
        """
        将Word文档转换为Markdown
        
        Args:
            word_path: Word文件路径
            output_path: 输出文件路径（可选）
        
        Returns:
            (markdown_content, success)
        """
        if not WORD_AVAILABLE:
            return "错误：python-docx库未安装。请运行: pip install python-docx", False
        
        try:
            doc = Document(word_path)
            markdown_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    markdown_content.append("\n")
                    continue
                
                # 根据样式判断标题级别
                style_name = para.style.name
                if 'Heading 1' in style_name or '标题 1' in style_name:
                    markdown_content.append(f"# {text}\n\n")
                elif 'Heading 2' in style_name or '标题 2' in style_name:
                    markdown_content.append(f"## {text}\n\n")
                elif 'Heading 3' in style_name or '标题 3' in style_name:
                    markdown_content.append(f"### {text}\n\n")
                elif 'Heading 4' in style_name or '标题 4' in style_name:
                    markdown_content.append(f"#### {text}\n\n")
                else:
                    markdown_content.append(f"{text}\n\n")
            
            # 处理表格
            for table in doc.tables:
                markdown_content.append("\n")
                # 表头
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                markdown_content.append("| " + " | ".join(headers) + " |\n")
                markdown_content.append("| " + " | ".join(["---"] * len(headers)) + " |\n")
                
                # 数据行
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_content.append("| " + " | ".join(cells) + " |\n")
                markdown_content.append("\n")
            
            result = "".join(markdown_content)
            
            # 保存到文件
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(result)
                logger.info(f"Word文档已转换为Markdown: {output_path}")
            
            return result, True
            
        except Exception as e:
            error_msg = f"Word转换失败: {str(e)}"
            logger.error(error_msg)
            return error_msg, False
    
    def convert_excel_to_markdown(self, excel_path: str, output_path: Optional[str] = None, sheet_name: Optional[str] = None) -> Tuple[str, bool]:
        """
        将Excel转换为Markdown
        
        Args:
            excel_path: Excel文件路径
            output_path: 输出文件路径（可选）
            sheet_name: 工作表名称（可选，如果为None则转换所有工作表）
        
        Returns:
            (markdown_content, success)
        """
        if not EXCEL_AVAILABLE:
            return "错误：pandas库未安装。请运行: pip install pandas openpyxl", False
        
        try:
            excel_file = pd.ExcelFile(excel_path)
            markdown_content = []
            
            sheets_to_process = [sheet_name] if sheet_name else excel_file.sheet_names
            
            for sheet in sheets_to_process:
                df = pd.read_excel(excel_path, sheet_name=sheet)
                
                markdown_content.append(f"## 工作表: {sheet}\n\n")
                
                # 转换为Markdown表格
                markdown_table = df.to_markdown(index=False)
                markdown_content.append(markdown_table)
                markdown_content.append("\n\n")
            
            result = "".join(markdown_content)
            
            # 保存到文件
            if output_path:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(result)
                logger.info(f"Excel已转换为Markdown: {output_path}")
            
            return result, True
            
        except Exception as e:
            error_msg = f"Excel转换失败: {str(e)}"
            logger.error(error_msg)
            return error_msg, False
    
    def convert(self, input_path: str, output_path: Optional[str] = None) -> Tuple[str, bool]:
        """
        自动识别文件类型并转换
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径（可选，如果不提供则自动生成）
        
        Returns:
            (markdown_content, success)
        """
        input_path = Path(input_path)
        
        if not input_path.exists():
            return f"错误：文件不存在: {input_path}", False
        
        # 自动生成输出路径
        if output_path is None:
            output_path = input_path.parent / f"{input_path.stem}.md"
        else:
            output_path = Path(output_path)
        
        # 根据文件扩展名选择转换方法
        ext = input_path.suffix.lower()
        
        if ext == '.pdf':
            return self.convert_pdf_to_markdown(str(input_path), str(output_path))
        elif ext in ['.docx', '.doc']:
            return self.convert_word_to_markdown(str(input_path), str(output_path))
        elif ext in ['.xlsx', '.xls']:
            return self.convert_excel_to_markdown(str(input_path), str(output_path))
        else:
            return f"错误：不支持的文件类型: {ext}", False
